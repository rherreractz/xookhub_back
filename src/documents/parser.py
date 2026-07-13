# Ruta: src/documents/parser.py
"""
Document parsing interface for the `documents` module.

Defines the contract `worker/tasks.py` uses to turn a raw uploaded file into
plain-text pages ready for chunking + embedding. Real extraction libraries
(`pypdf`, `python-docx`, `unstructured`, ...) get wired in behind this same
interface later — nothing upstream of this module should need to change
when that happens, only the bodies of the `DocumentParser` subclasses below.
"""

from __future__ import annotations

from abc import ABC, abstractmethod
from dataclasses import dataclass, field


@dataclass(frozen=True, slots=True)
class ParsedPage:
    """A single logical page/unit of extracted text.

    `page_number` is `None` for formats without a native pagination concept
    (plain text, markdown) — chunks derived from it simply carry no page
    reference downstream.
    """

    page_number: int | None
    content: str


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    """Full extraction result for one uploaded file — the output contract
    every `DocumentParser.parse()` implementation must return."""

    pages: list[ParsedPage]
    metadata: dict[str, str] = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.content for page in self.pages if page.content)


class DocumentParser(ABC):
    """Strategy interface — one implementation per supported mime type."""

    supported_mime_types: tuple[str, ...] = ()

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """Extract text from `file_path` (a local path — the caller is
        responsible for downloading the object from MinIO first) and
        return it as a `ParsedDocument`."""
        raise NotImplementedError


class PlainTextParser(DocumentParser):
    """Real implementation — plain text needs no extraction library."""

    supported_mime_types = ("text/plain", "text/markdown")

    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, encoding="utf-8", errors="replace") as fh:
            content = fh.read()
        return ParsedDocument(pages=[ParsedPage(page_number=1, content=content)])


class PDFParser(DocumentParser):
    """Real extraction via `pypdf`.

    `page.extract_text()` can legitimately return `None`/empty for a page
    with no extractable text layer (e.g. a scanned image with no OCR) —
    that page just contributes an empty `ParsedPage.content` rather than
    raising, since a PDF with some real pages and one blank scanned cover
    page shouldn't fail the whole document.
    """

    supported_mime_types = ("application/pdf",)

    def parse(self, file_path: str) -> ParsedDocument:
        import pypdf

        reader = pypdf.PdfReader(file_path)
        pages = [
            ParsedPage(page_number=i + 1, content=page.extract_text() or "")
            for i, page in enumerate(reader.pages)
        ]
        return ParsedDocument(
            pages=pages,
            metadata={"parser": "pypdf", "num_pages": str(len(pages))},
        )


class DocxParser(DocumentParser):
    """Real extraction via `python-docx`.

    .docx has no native page concept in the file format itself (pagination
    is a rendering-time computation, not stored data), so this stays a
    single `ParsedPage` with `page_number=None`, matching `PlainTextParser`.
    """

    supported_mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def parse(self, file_path: str) -> ParsedDocument:
        import docx

        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return ParsedDocument(
            pages=[ParsedPage(page_number=None, content=text)],
            metadata={"parser": "python-docx"},
        )


class CodeParser(DocumentParser):
    """Source code files — plain text extraction, same as
    `PlainTextParser`, but registered under the canonical `text/x-code`
    mime type that `DocumentService.upload()` normalizes every accepted
    code extension to (see the note there on why: browsers guess
    Content-Type for code files inconsistently, so extension — not
    Content-Type — is the actual signal for "this is code")."""

    supported_mime_types = ("text/x-code",)

    def parse(self, file_path: str) -> ParsedDocument:
        with open(file_path, encoding="utf-8", errors="replace") as fh:
            content = fh.read()
        return ParsedDocument(pages=[ParsedPage(page_number=1, content=content)])


_PARSERS: tuple[DocumentParser, ...] = (
    PlainTextParser(),
    PDFParser(),
    DocxParser(),
    CodeParser(),
)


class UnsupportedMimeTypeError(ValueError):
    """Raised when no registered `DocumentParser` handles a given mime type."""


def get_parser_for(mime_type: str) -> DocumentParser:
    for parser in _PARSERS:
        if mime_type in parser.supported_mime_types:
            return parser
    raise UnsupportedMimeTypeError(
        f"No hay parser registrado para mime_type={mime_type!r}"
    )


def parse_document(file_path: str, mime_type: str) -> ParsedDocument:
    """Entry point used by `worker.tasks.process_document_task`."""
    return get_parser_for(mime_type).parse(file_path)


def chunk_text(text: str, *, chunk_size: int = 1_000, overlap: int = 150) -> list[str]:
    """Naive fixed-size character chunker with overlap.

    Good enough as a first pass for `document_chunks.content`; swap for a
    token-aware / sentence-boundary-aware splitter (e.g. `tiktoken`-based)
    once the `rag` module lands and chunk quality starts to matter for
    retrieval relevance.
    """
    if chunk_size <= overlap:
        raise ValueError("chunk_size debe ser mayor que overlap.")

    text = text.strip()
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = start + chunk_size
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= text_len:
            break
        start = end - overlap
    return chunks